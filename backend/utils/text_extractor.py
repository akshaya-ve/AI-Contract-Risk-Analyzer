"""
Document text extraction utilities.

Supports:
  - PDF  via PyMuPDF (fitz)  — preserves page numbers
  - DOCX via python-docx     — extracts paragraphs and table cells

Design principles:
  - Every extracted chunk carries its page/section number so the RAG pipeline
    can surface accurate source references in the chat UI.
  - Raises typed exceptions from utils.exceptions so the API layer can map
    them to correct HTTP status codes without any try/except at the route level.
  - All public functions are pure (no side-effects) — easy to unit-test.
"""

import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from backend.utils.exceptions import (
    DocumentExtractionError,
    EmptyDocumentError,
    UnsupportedFileTypeError,
)
from backend.utils.logger import get_logger

logger = get_logger(__name__)

# ─── Data Structures ──────────────────────────────────────────────────────────


@dataclass
class PageContent:
    """
    A single page / section of extracted text.

    Attributes:
        page_number: 1-indexed page number (PDF) or paragraph block index (DOCX).
        text:        Raw extracted text for this page/section.
        char_count:  Character count (useful for filtering near-empty pages).
    """

    page_number: int
    text: str
    char_count: int = field(init=False)

    def __post_init__(self) -> None:
        self.char_count = len(self.text.strip())


@dataclass
class ExtractionResult:
    """
    Full extraction result for an uploaded document.

    Attributes:
        filename:       Original file name.
        file_type:      ".pdf" or ".docx".
        pages:          Ordered list of PageContent objects.
        full_text:      Concatenated text of all pages (used for chunking).
        total_pages:    Total number of pages/sections extracted.
        total_chars:    Total character count across the document.
    """

    filename: str
    file_type: str
    pages: list[PageContent]
    full_text: str
    total_pages: int
    total_chars: int


# ─── PDF Extraction ───────────────────────────────────────────────────────────


def extract_pdf(file_path: str | Path) -> ExtractionResult:
    """
    Extract text from a PDF file using PyMuPDF (fitz).

    Iterates over every page and extracts text blocks, preserving page numbers
    so the RAG pipeline can return accurate source citations.

    Args:
        file_path: Absolute path to the PDF file.

    Returns:
        ExtractionResult with per-page content.

    Raises:
        DocumentExtractionError: If PyMuPDF fails to open or parse the file.
        EmptyDocumentError:      If the PDF contains no extractable text
                                 (e.g. scanned image-only PDF).
    """
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise DocumentExtractionError(
            str(file_path), "PyMuPDF (fitz) is not installed."
        ) from e

    file_path = Path(file_path)
    filename = file_path.name
    logger.info(f"Extracting PDF: {filename}")

    try:
        doc = fitz.open(str(file_path))
    except Exception as e:
        raise DocumentExtractionError(filename, str(e)) from e

    pages: list[PageContent] = []

    try:
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            raw_text = page.get_text("text")  # type: ignore[attr-defined]

            # Normalise whitespace: collapse 3+ consecutive newlines to 2
            cleaned = _clean_text(raw_text)

            if cleaned:  # skip genuinely blank pages
                pages.append(PageContent(page_number=page_num + 1, text=cleaned))

    except Exception as e:
        raise DocumentExtractionError(filename, f"Page extraction failed: {e}") from e
    finally:
        doc.close()

    if not pages:
        raise EmptyDocumentError(filename)

    full_text = _join_pages(pages)
    total_chars = sum(p.char_count for p in pages)

    logger.info(
        f"PDF extracted: {filename} | pages={len(pages)} | chars={total_chars:,}"
    )

    return ExtractionResult(
        filename=filename,
        file_type=".pdf",
        pages=pages,
        full_text=full_text,
        total_pages=len(pages),
        total_chars=total_chars,
    )


# ─── DOCX Extraction ─────────────────────────────────────────────────────────


def extract_docx(file_path: str | Path) -> ExtractionResult:
    """
    Extract text from a DOCX file using python-docx.

    Extracts:
      - All paragraph text (including headings, body, list items).
      - All table cells (row by row, cell by cell).

    Paragraphs are grouped into logical "pages" of ~50 paragraphs each
    so the RAG pipeline has meaningful page_number metadata.

    Args:
        file_path: Absolute path to the DOCX file.

    Returns:
        ExtractionResult with per-section content.

    Raises:
        DocumentExtractionError: If python-docx fails to parse the file.
        EmptyDocumentError:      If no text is found.
    """
    try:
        import docx  # python-docx
    except ImportError as e:
        raise DocumentExtractionError(
            str(file_path), "python-docx is not installed."
        ) from e

    file_path = Path(file_path)
    filename = file_path.name
    logger.info(f"Extracting DOCX: {filename}")

    try:
        document = docx.Document(str(file_path))
    except Exception as e:
        raise DocumentExtractionError(filename, str(e)) from e

    raw_blocks: list[str] = []

    # ── Paragraphs ──────────────────────────────────────────────────────────
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            raw_blocks.append(text)

    # ── Tables ───────────────────────────────────────────────────────────────
    for table in document.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                raw_blocks.append(" | ".join(row_cells))

    if not raw_blocks:
        raise EmptyDocumentError(filename)

    # Group blocks into virtual "pages" of 50 blocks each
    chunk_size = 50
    pages: list[PageContent] = []
    for idx in range(0, len(raw_blocks), chunk_size):
        group = raw_blocks[idx : idx + chunk_size]
        page_text = _clean_text("\n".join(group))
        if page_text:
            pages.append(
                PageContent(page_number=(idx // chunk_size) + 1, text=page_text)
            )

    if not pages:
        raise EmptyDocumentError(filename)

    full_text = _join_pages(pages)
    total_chars = sum(p.char_count for p in pages)

    logger.info(
        f"DOCX extracted: {filename} | sections={len(pages)} | chars={total_chars:,}"
    )

    return ExtractionResult(
        filename=filename,
        file_type=".docx",
        pages=pages,
        full_text=full_text,
        total_pages=len(pages),
        total_chars=total_chars,
    )


# ─── Public Dispatcher ────────────────────────────────────────────────────────


def extract_document(file_path: str | Path) -> ExtractionResult:
    """
    Auto-detect file type and dispatch to the appropriate extractor.

    This is the single entry point used by DocumentService.

    Args:
        file_path: Absolute path to the uploaded file.

    Returns:
        ExtractionResult populated by pdf or docx extractor.

    Raises:
        UnsupportedFileTypeError: If the extension is not .pdf or .docx.
        DocumentExtractionError:  Propagated from the specific extractor.
        EmptyDocumentError:       Propagated from the specific extractor.
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return extract_pdf(file_path)
    elif suffix == ".docx":
        return extract_docx(file_path)
    else:
        raise UnsupportedFileTypeError(file_path.name)


# ─── File Validation ─────────────────────────────────────────────────────────


def validate_file(
    filename: str,
    file_size_bytes: int,
    max_size_mb: int = 20,
    allowed_extensions: Optional[list[str]] = None,
) -> None:
    """
    Validate uploaded file before saving to disk.

    Checks:
      1. File extension is in the allowed list.
      2. File size does not exceed the configured maximum.

    Args:
        filename:           Original filename from the upload.
        file_size_bytes:    File size in bytes.
        max_size_mb:        Maximum allowed size in megabytes.
        allowed_extensions: List of allowed extensions e.g. [".pdf", ".docx"].

    Raises:
        UnsupportedFileTypeError: If the extension is not allowed.
        DocumentExtractionError:  If the file exceeds the size limit.
    """
    if allowed_extensions is None:
        allowed_extensions = [".pdf", ".docx"]

    suffix = Path(filename).suffix.lower()

    if suffix not in allowed_extensions:
        raise UnsupportedFileTypeError(filename)

    max_bytes = max_size_mb * 1024 * 1024
    if file_size_bytes > max_bytes:
        raise DocumentExtractionError(
            filename,
            f"File size {file_size_bytes / (1024*1024):.1f} MB exceeds "
            f"the {max_size_mb} MB limit.",
        )

    logger.debug(
        f"File validated: {filename} "
        f"({file_size_bytes / 1024:.1f} KB, type={suffix})"
    )


# ─── Private Helpers ─────────────────────────────────────────────────────────


def _clean_text(text: str) -> str:
    """
    Normalise extracted text:
      - Strip leading/trailing whitespace per line.
      - Collapse 3+ consecutive blank lines into 2 (preserve paragraph breaks).
      - Remove null bytes and other control characters.
    """
    import re

    # Remove null bytes / control chars (except newlines and tabs)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

    # Collapse excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def _join_pages(pages: list[PageContent]) -> str:
    """
    Join page contents with a clear page separator so the chunker
    can optionally split on page boundaries.
    """
    parts: list[str] = []
    for page in pages:
        parts.append(f"[Page {page.page_number}]\n{page.text}")
    return "\n\n".join(parts)
